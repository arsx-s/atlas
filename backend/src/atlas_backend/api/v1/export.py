"""Export API routes."""

from __future__ import annotations

import json
from typing import Any

from atlas_backend.core.errors import AtlasErrorCode, AtlasException
from atlas_backend.core.responses import create_success_response
from atlas_backend.persistence import AtlasLocalStore


def _build_report_data(store: AtlasLocalStore, session_id: str) -> dict[str, Any]:
    session = store.get_research_session(session_id)
    if not session:
        raise AtlasException(AtlasErrorCode.NOT_FOUND, "Research session not found.")
    report = store.get_report(session_id)
    evidence = store.list_evidence(session_id)
    return {"session": session, "report": report, "evidence": evidence}


def _render_markdown(data: dict[str, Any]) -> str:
    session = data["session"]
    report = data["report"]
    evidence = data["evidence"]
    lines = [f"# {session['title']}", "", f"**Query:** {session['query']}", ""]
    if report:
        lines.append("## Executive Summary")
        lines.append(report.get("executive_summary", ""))
        lines.append("")
        for section in report.get("sections", []):
            lines.append(f"## {section.get('heading', 'Section')}")
            lines.append(section.get("content", ""))
            lines.append("")
    else:
        lines.append("*No report generated yet.*")
        lines.append("")
    if evidence:
        lines.append("## Evidence")
        lines.append("")
        for e in evidence:
            source = e.get("source_url") or e.get("source_document") or ""
            lines.append(f"- **{e['title']}** (confidence: {e.get('confidence_score', 1.0)})")
            if source:
                lines.append(f"  - Source: {source}")
            lines.append("")
    if report and report.get("citations"):
        lines.append("## Citations")
        lines.append("")
        for i, c in enumerate(report["citations"], 1):
            lines.append(f"[{i}] {c.get('text', '')} — {c.get('source', '')}")
    return "\n".join(lines)


def _markdown_to_html(md_text: str) -> str:
    try:
        import markdown
        html_body = markdown.markdown(md_text, extensions=["extra"])
    except ImportError:
        html_body = f"<pre>{md_text}</pre>"
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8"></head><body>{html_body}</body></html>"""


def create_export_router(store: AtlasLocalStore) -> Any:
    from fastapi import APIRouter
    from fastapi.responses import PlainTextResponse

    router = APIRouter(tags=["export"])

    @router.get("/export/{session_id}/markdown")
    async def export_markdown(session_id: str) -> Any:
        data = _build_report_data(store, session_id)
        content = _render_markdown(data)
        return PlainTextResponse(
            content, media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=atlas-report-{session_id[:8]}.md"},
        )

    @router.get("/export/{session_id}/json")
    async def export_json(session_id: str) -> Any:
        data = _build_report_data(store, session_id)
        from fastapi.responses import JSONResponse
        return JSONResponse(
            content=create_success_response(data).model_dump(mode="json"),
            headers={"Content-Disposition": f"attachment; filename=atlas-report-{session_id[:8]}.json"},
        )

    @router.get("/export/{session_id}/pdf")
    async def export_pdf(session_id: str) -> Any:
        data = _build_report_data(store, session_id)
        md = _render_markdown(data)
        html = _markdown_to_html(md)
        try:
            from weasyprint import HTML
            pdf_bytes = HTML(string=html).write_pdf()
        except ImportError:
            raise AtlasException(AtlasErrorCode.INTERNAL_ERROR, "PDF export requires weasyprint.")
        from fastapi.responses import Response
        return Response(
            content=pdf_bytes, media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=atlas-report-{session_id[:8]}.pdf"},
        )

    @router.get("/export/{session_id}/docx")
    async def export_docx(session_id: str) -> Any:
        data = _build_report_data(store, session_id)
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            raise AtlasException(AtlasErrorCode.INTERNAL_ERROR, "DOCX export requires python-docx.")
        doc = Document()
        session = data["session"]
        report = data["report"]
        evidence = data["evidence"]
        doc.add_heading(session["title"], level=1)
        doc.add_paragraph(f"Query: {session['query']}")
        if report:
            doc.add_heading("Executive Summary", level=2)
            doc.add_paragraph(report.get("executive_summary", ""))
            for section in report.get("sections", []):
                doc.add_heading(section.get("heading", "Section"), level=2)
                doc.add_paragraph(section.get("content", ""))
        else:
            doc.add_paragraph("No report generated yet.")
        if evidence:
            doc.add_heading("Evidence", level=2)
            for e in evidence:
                p = doc.add_paragraph()
                run = p.add_run(e["title"])
                run.bold = True
                p.add_run(f" (confidence: {e.get('confidence_score', 1.0)})")
                source = e.get("source_url") or e.get("source_document") or ""
                if source:
                    p.add_run(f"\nSource: {source}")
        if report and report.get("citations"):
            doc.add_heading("Citations", level=2)
            for i, c in enumerate(report["citations"], 1):
                doc.add_paragraph(f"[{i}] {c.get('text', '')} — {c.get('source', '')}")
        import io
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        from fastapi.responses import Response
        return Response(
            content=buf.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=atlas-report-{session_id[:8]}.docx"},
        )

    return router
